from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path("C:/Users/dell/Downloads/\u4e2d\u671f\u62a5\u544a-\u80e1\u6768\u8d85.docx")
OUT = Path("C:/Users/dell/Downloads/\u4e2d\u671f\u62a5\u544a-\u80e1\u6768\u8d85-\u4fee\u6539\u5b8c\u5584\u7248.docx")


def set_run_font(run, east_asia: str = "宋体", size: float = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), "Times New Roman")


def set_paragraph_format(
    paragraph,
    *,
    first_line: bool = False,
    align=None,
    line_spacing: float = 1.25,
    space_after: float = 0,
) -> None:
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(24) if first_line else Pt(0)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if align is not None:
        paragraph.alignment = align


def clear_cell(cell) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_para(
    cell,
    text: str = "",
    *,
    style: str = "body",
    size: float = 12,
    first_line: bool | None = None,
    space_after: float = 0,
) :
    paragraph = cell.add_paragraph()
    if style == "section":
        run = paragraph.add_run(text)
        set_run_font(run, "黑体", size, bold=True)
        set_paragraph_format(paragraph, first_line=False, line_spacing=1.2, space_after=2)
    elif style == "subsection":
        run = paragraph.add_run(text)
        set_run_font(run, "黑体", size, bold=True)
        set_paragraph_format(paragraph, first_line=False, line_spacing=1.2, space_after=1)
    elif style == "prompt":
        run = paragraph.add_run(text)
        set_run_font(run, "楷体", size, bold=False)
        set_paragraph_format(paragraph, first_line=False, line_spacing=1.2, space_after=2)
    elif style == "caption":
        run = paragraph.add_run(text)
        set_run_font(run, "宋体", size, bold=True)
        set_paragraph_format(paragraph, first_line=False, line_spacing=1.2, space_after=1)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, "宋体", size, bold=False)
        if first_line is None:
            first_line = True
        set_paragraph_format(paragraph, first_line=first_line, line_spacing=1.25, space_after=0)
    return paragraph


def replace_paragraph_text(paragraph, text: str, east_asia: str, size: float, bold: bool = False, align=None) -> None:
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia, size, bold)
    if align is not None:
        paragraph.alignment = align


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def normalize_table_body(table) -> None:
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            set_cell_margins(cell)


MAIN_PROGRESS = [
    ("section", "研究目标"),
    ("body", "本课题面向复杂工程优化中计算代价高、目标函数黑箱、多峰搜索与梯度信息缺失等问题，围绕“获得性遗传算法（Hybrid Algorithm, HA）”开展研究。当前工作以聚类小生境、局部搜索、代理模型与混合遗传算子为核心，目标是在有限函数评估次数下提升全局搜索效率、局部开发能力和约束处理能力，并进一步扩展到多目标优化场景。"),
    ("body", "现阶段代码库 HA 已形成较完整的单目标与多目标算法实现。其中，ha_Nelder_Mead.py 负责单目标 HA 主体，ha_nsga3.py 负责多目标 HA-NSGA-III 扩展，两者均可接入 pymoo 优化流程，支持重复实验、结果汇总与可视化分析。"),
    ("section", "已完成的主要研究内容"),
    ("subsection", "1. 单目标 HA 算法框架实现"),
    ("body", "完成了单目标 HA 的工程化实现，核心类 HA 继承 pymoo.core.algorithm.Algorithm，包含种群初始化、批量评价、约束违反度计算、历史评估缓存、精英保留、后代生成、变异、去重和环境选择等流程。PopulationHistory 类用于记录已评价解，按决策变量容差去重并缓存目标值与约束违反度，减少局部搜索中的重复函数评估。"),
    ("body", "算法采用约束优先排序规则：可行解优先于不可行解，可行解之间按目标函数值排序，不可行解之间按约束违反度排序。该机制使算法能够统一处理无约束、有边界约束和一般不等式约束问题，为后续工程仿真优化提供了基础。"),
    ("subsection", "2. 聚类小生境与触发式局部学习机制"),
    ("body", "在单目标 HA 中实现了 K-means、MeanShift 和 DBSCAN 三类聚类方法。K-means 适合预设小生境数量的场景；MeanShift 可依据数据密度自动确定簇数；DBSCAN 能发现非球形簇，并支持 eps 与 min_samples 的自动估计。"),
    ("body", "局部学习并非每代固定执行，而是采用“周期触发 + 停滞触发”的策略：每 5 代周期性触发一次，或在连续若干代没有改进时触发。该设计避免局部搜索过度消耗函数评估次数，同时在算法停滞时增强跳出局部最优的能力。"),
    ("subsection", "3. 多策略局部搜索与代理辅助优化"),
    ("body", "围绕黑箱优化和梯度不可得问题，已实现 Nelder-Mead、RBF、GP、history-ladder、L-BFGS-B、TNC、SLSQP、Powell、trust-constr，以及 Adam、AdamW、Lion、Sophia 等局部搜索接口。其中 Nelder-Mead 版本利用历史点构造单纯形，RBF 和 GP 版本利用历史评估数据构建局部代理模型，再在代理模型上执行 L-BFGS-B 搜索。"),
    ("body", "RBF 局部搜索使用近邻历史点拟合径向基函数代理模型，GP 局部搜索采用高斯过程与置信下界思想兼顾预测均值和不确定性。两类方法均通过历史缓存减少真实函数调用，适合后续与 SolidWorks 或其他昂贵仿真模型结合。"),
    ("subsection", "4. 混合遗传算子与多样性维护"),
    ("body", "已完成三类遗传操作的组合：适应度加权 SBX 交叉、多父代加权交叉和 DE/current-to-best 差分进化策略。三者分别强调稳定重组、多优秀个体信息融合和向优良区域定向推进，使 HA 在探索与开发之间保持平衡。"),
    ("body", "变异阶段采用多项式变异和方向性/边界变异等策略，并在种群严重趋同时注入一定比例随机个体、重置搜索步长，以缓解早熟收敛和后期多样性不足问题。"),
    ("subsection", "5. 多目标 HA-NSGA-III 扩展"),
    ("body", "在 ha_nsga3.py 中完成多目标版本 HA_NSGA3。该版本继承单目标 HA 框架，引入 MOPopulationHistory 支持多维目标值缓存，并使用非支配排序、NSGA-III 参考方向关联和 niching 环境选择替代单目标适应度排序。"),
    ("body", "多目标局部搜索方面，已实现 PBI、Tchebycheff、Weighted Sum、ASF、AASF 和 IPBI 等标量化方式，将单目标局部搜索器适配到多目标场景；同时实现了 MO-Nelder-Mead 支配关系版本，直接依据可行性、Pareto 支配关系和到参考方向的垂直距离进行单纯形顶点比较。"),
    ("body", "针对多目标前沿分布性，已加入动态 nadir 估计、PBI 惩罚系数调整和覆盖率触发的定向注入机制，并通过 niche_strategy 在参考方向小生境与 K-means 小生境之间切换，用于后续消融实验。"),
    ("subsection", "6. 阶段性实验与结果整理"),
    ("body", "单目标方面，已在 F2、F3、F4、Ackley、Griewank 等测试函数上形成 HA 与 GA 及不同聚类策略的对比结果，结果文件保存在 results、experiment_results 和 experiment_results_nelder_mead 等目录中，并配套生成收敛曲线与指标图。"),
    ("body", "多目标方面，已完成 NSHA 与 pymoo 内置 NSGA-III 在 ZDT1、ZDT2、ZDT3、ZDT4、ZDT6 上的初步对比。当前 summary 结果表明，NSHA 的部分改进配置在 ZDT2、ZDT6 等问题上取得较低 IGD；在 ZDT4 这类多峰问题上，NSHA-Base 相比 NSGA-III 表现出更强的局部开发能力。但不同问题上的优势并不完全一致，仍需通过更多种子、更多指标和参数消融进一步验证。"),
    ("section", "与开题计划的对应情况"),
    ("body", "总体来看，课题已完成从单目标 HA 到多目标 HA-NSGA-III 的核心代码实现，基本覆盖开题计划中“聚类局部搜索、混合交叉、约束处理、多目标扩展和基准实验验证”的主要任务。后续工作将重点从算法功能实现转向实验系统化、参数敏感性分析、工程仿真案例验证和论文写作。"),
]


ACHIEVEMENTS = [
    ("prompt", "按《研究生学位论文撰写格式规范》的格式要求分类填写与学位论文相关的阶段性研究成果，例如期刊论文、会议论文、科研获奖、专利、制定标准等，限填第一作者或导师为第一作者时的第二作者成果，其中已录用、已投稿或拟投稿的在括号内注明（可续页）"),
    ("caption", "阶段性研究成果"),
    ("body", "1. 代码成果：完成单目标 HA 核心实现 ha_Nelder_Mead.py 和多目标 HA-NSGA-III 实现 ha_nsga3.py，形成可运行、可复现实验的算法代码库。"),
    ("body", "2. 实验成果：完成单目标基准函数、不同聚类策略、局部搜索策略以及多目标 ZDT 系列问题的阶段性对比实验，形成 CSV 汇总、收敛曲线、Pareto 前沿分布图和对比图。"),
    ("body", "3. 文档成果：已整理 README、README_HA、指标说明等说明文件，对算法结构、运行方式、局部搜索效率指标和结果目录进行了规范化记录。"),
    ("body", "4. 论文成果：目前尚未形成已录用或已发表论文。下一阶段拟在完成补充实验和统计检验后，围绕“聚类小生境代理辅助混合进化算法”和“参考方向驱动的多目标 HA-NSGA-III”整理投稿论文或学位论文核心章节。"),
]


PROBLEMS = [
    ("prompt", "1. 未按开题计划完成的研究工作，研究工作存在的原理性、技术性难题以及在实验条件等方面的限制（可续页）"),
    ("caption", "存在的主要问题"),
    ("body", "1. 多目标算法的稳定性仍需增强。现有实验表明，动态 nadir、PBI 参数和覆盖率注入在不同 ZDT 问题上的效果存在差异，说明多目标小生境分布性与局部搜索开发强度之间仍需进一步平衡。"),
    ("body", "2. 局部搜索与函数评估预算之间的关系仍需量化。Nelder-Mead、RBF、GP 等局部搜索能够提升局部开发能力，但在昂贵黑箱问题中会增加函数评估成本，需要明确不同阶段、不同个体类型下的触发条件和迭代预算。"),
    ("body", "3. 实验统计规模仍需扩大。目前已有多组基准实验和可视化结果，但还需要增加随机种子、对比算法和统计检验，避免结论依赖单次运行或个别问题特性。"),
    ("body", "4. 工程仿真案例仍需进一步闭环验证。代码库已具备 QuickSimu 和代理模型相关实验基础，但与真实工程仿真流程的输入输出、约束定义、计算耗时和代理模型误差传递仍需系统梳理。"),
    ("body", "5. 理论分析尚不充分。目前主要完成算法实现和实验验证，关于获得性遗传机制、混合交叉算子、局部搜索触发机制对收敛性的影响，还需进一步凝练理论说明。"),
]


SOLUTIONS = [
    ("prompt", "2. 针对上述问题采取何种解决办法，对学位论文的研究内容及所采取的理论方法、技术路线和实施方案的进一步调整，以及下一步的研究计划（可续页）"),
    ("caption", "解决办法与下一步研究计划"),
    ("body", "1. 完善多目标 HA-NSGA-III 的对比实验。继续以 ZDT、DTLZ 等标准测试问题为基础，补充 IGD、HV、Spacing、Spread 等指标，并与 NSGA-II、NSGA-III、MOEA/D 等典型算法进行多随机种子统计对比。"),
    ("body", "2. 开展关键模块消融实验。围绕动态 nadir 估计、PBI 惩罚系数、覆盖率注入、参考方向小生境、K-means 小生境、RBF/GP/Nelder-Mead 局部搜索等模块，设计单因素和组合消融，明确各模块对收敛性、分布性和函数评估成本的影响。"),
    ("body", "3. 优化局部搜索触发与预算分配。根据停滞检测、历史改进率、局部搜索效率和前沿覆盖率动态调整局部搜索深度，避免在低收益阶段消耗过多真实函数评估次数。"),
    ("body", "4. 推进工程仿真与代理模型验证。以 QuickSimu1、QuickSimu2 及后续 SolidWorks 结构优化案例为对象，比较 RBF、GP、Kriging、Polynomial、KAN/KAN-GP 等代理模型的拟合精度、优化效果和真实回评表现。"),
    ("body", "5. 完善理论分析和论文撰写。围绕混合进化框架、聚类小生境、局部搜索接受准则、约束处理和多目标参考方向机制，形成学位论文的算法章节、实验章节和理论讨论。"),
    ("caption", "后续可能的工作标题"),
    ("body", "（1）面向昂贵黑箱优化的聚类小生境代理辅助混合进化算法研究。"),
    ("body", "（2）基于参考方向小生境与支配关系的多目标 HA-NSGA-III 算法研究。"),
    ("body", "（3）动态 nadir 估计与覆盖率注入机制对多目标混合进化算法分布性的影响分析。"),
    ("body", "（4）复杂工程仿真场景下获得性遗传算法的代理模型加速与应用验证。"),
]


def fill_cell(cell, content) -> None:
    clear_cell(cell)
    for style, text in content:
        add_para(cell, text, style=style)


def add_plain_paragraph(cell, text: str, *, size: float = 12, bold: bool = False, align=None):
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", size, bold)
    set_paragraph_format(paragraph, first_line=False, align=align, line_spacing=1.15, space_after=0)
    return paragraph


def add_blank_lines(cell, count: int, *, size: float = 12) -> None:
    for _ in range(count):
        add_plain_paragraph(cell, "", size=size)


def fix_review_table(table) -> None:
    """Clean up signature/date lines so the school template does not wrap dates."""
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            set_cell_margins(cell, top=90, start=110, bottom=90, end=110)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 1. 导师意见
    cell = table.cell(0, 0)
    clear_cell(cell)
    add_plain_paragraph(cell, "1.导师对工作进展及研究计划的意见：", bold=False)
    add_blank_lines(cell, 4)
    add_plain_paragraph(cell, "校内导师（组）签字：____________________    年    月    日", size=11)
    add_plain_paragraph(cell, "校外导师签字：________________________    年    月    日", size=11)

    # 2. 中期考评专家组意见标题
    cell = table.cell(1, 0)
    clear_cell(cell)
    add_plain_paragraph(cell, "2.中期考评专家组意见", bold=True)

    # Keep compact label cells centered.
    for pos in [(2, 0), (2, 2), (3, 0), (4, 0), (5, 0)]:
        for paragraph in table.cell(*pos).paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 结论说明重写，避免挤压。
    cell = table.cell(5, 1)
    clear_cell(cell)
    add_plain_paragraph(cell, "□通过              □原则通过              □不通过", size=12)
    add_plain_paragraph(cell, "通过：表决票均为合格", size=9, bold=True)
    add_plain_paragraph(cell, "原则通过：表决票中有 1 票为基本合格或不合格，其余为合格和基本合格", size=9)
    add_plain_paragraph(cell, "不通过：表决票中有 2 票及以上为不合格", size=9)

    # 专家组建议与签名。
    cell = table.cell(6, 0)
    clear_cell(cell)
    add_plain_paragraph(cell, "对学位论文工作进展以及下一步研究计划的建议，是否适合继续攻读学位：")
    add_blank_lines(cell, 5)
    add_plain_paragraph(cell, "专家组签名：____________________________    年    月    日", size=11)

    # 学院意见与负责人签名。
    cell = table.cell(7, 0)
    clear_cell(cell)
    add_plain_paragraph(cell, "3.学院意见：", bold=True)
    add_blank_lines(cell, 5)
    add_plain_paragraph(cell, "负责人签名：____________________________    年    月    日", size=11)

    # The original last row only held a wrapped date fragment. Leave it blank and shallow.
    cell = table.cell(8, 0)
    clear_cell(cell)
    add_plain_paragraph(cell, "", size=6)
    table.rows[8].height = Pt(6)


def main() -> None:
    doc = Document(str(SRC))

    # Cover page corrections.
    replace_paragraph_text(
        doc.paragraphs[1],
        "专业学位研究生学位论文中期报告表",
        "黑体",
        23,
        False,
        WD_ALIGN_PARAGRAPH.CENTER,
    )
    replace_paragraph_text(
        doc.paragraphs[8],
        "\t论文题目： 基于聚类局部搜索与混合交叉的",
        "黑体",
        15,
        False,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    replace_paragraph_text(
        doc.paragraphs[9],
        "\t               获得性遗传算法研究与应用",
        "黑体",
        15,
        False,
        None,
    )
    replace_paragraph_text(
        doc.paragraphs[12],
        "\t填表日期：          2026 年  6 月 30 日",
        "黑体",
        15,
        False,
        None,
    )

    # Normalize body heading paragraphs.
    for idx in (15, 16, 17):
        if idx < len(doc.paragraphs) and doc.paragraphs[idx].text.strip():
            replace_paragraph_text(
                doc.paragraphs[idx],
                doc.paragraphs[idx].text.strip(),
                "黑体",
                14,
                False,
                None,
            )

    for table in doc.tables:
        normalize_table_body(table)

    # Main report form content.
    fill_cell(doc.tables[0].cell(4, 0), MAIN_PROGRESS)
    fill_cell(doc.tables[0].cell(6, 0), ACHIEVEMENTS)
    fill_cell(doc.tables[1].cell(0, 0), PROBLEMS)
    fill_cell(doc.tables[1].cell(1, 0), SOLUTIONS)

    # Keep review/opinion areas blank, only normalize their fonts where template text exists.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            size = run.font.size.pt if run.font.size else 12
                            # Preserve cover/table prompt boldness, normalize Chinese fallback.
                            east_asia = "黑体" if run.font.bold else "宋体"
                            if "导师" in run.text or "考评" in run.text or "学院" in run.text:
                                east_asia = "宋体"
                            set_run_font(run, east_asia=east_asia, size=size, bold=bool(run.font.bold))

    fix_review_table(doc.tables[2])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(str(OUT))


if __name__ == "__main__":
    main()
